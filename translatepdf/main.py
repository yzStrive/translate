#coding:utf-8
# pip install python-docx
# or from https://pypi.python.org/pypi/python-docx/0.8.6 download it then excute python setup.py install
import translate_by_google
import os
import sys
import parsepdf
from docx import Document
import docx
cur_path = os.path.split(os.path.realpath(sys.argv[0]))[0]


def main(input_file_path,file_name=''):
    # 将pdf转换成str数组，结果中每一项为每一段
    results = parsepdf.parse(input_file_path)
    fp_e_c_document, fp_o_c_document = Document(), Document()
    # 每段的分隔符
    flag = '************************************'
    for item in results:
        # 每一段是'\n'来连接的，需要处理下
        item = ''.join(item.split('\n'))
        if len(item) <= 40:
            fp_e_c_document.add_paragraph(flag)
            fp_o_c_document.add_paragraph(flag)
        else:
            translate_r = translate_by_google.translate_google(item)
            translate_r = handle_response_from_translate(translate_r)
            fp_e_c_document.add_paragraph(translate_r['e_c'])
            fp_e_c_document.add_paragraph(flag)

            fp_o_c_document.add_paragraph(translate_r['o_c'])
            fp_o_c_document.add_paragraph(flag)
    fp_e_c_document.add_page_break()
    fp_o_c_document.add_page_break()
    # 输出word路径
    fp_e_c_document.save(cur_path + '/output/' + file_name + '_ec.docx')
    fp_o_c_document.save(cur_path + '/output/' + file_name + '_oc.docx')


def handle_response_from_translate(cont):
    # google翻译的结构需要处理才能得到想要的
    only_chinese = ""
    chinese_english = ""
    for item in cont[0]:
        if item[0] and item[1]:
            translate_result = item[0]
            translate_source = item[1]
            only_chinese = only_chinese + translate_result
            chinese_english = chinese_english + \
                translate_source + '\n' + translate_result + '\n'
    return {'e_c': chinese_english, 'o_c': only_chinese}


if __name__ == '__main__':
    # 待翻译的pdf路径
    input_suffix = '.pdf'
    input_file_name = 'Spatial_dynamic_patterns_of_hand-foot-mouth_diseas'
    path = cur_path + '/input/' + input_file_name + input_suffix
    main(path,input_file_name)
